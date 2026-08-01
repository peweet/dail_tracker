"""
pq_ida_land_tables.py — EXPERIMENTAL: extract the <table> markup that
pq_answer_mine_experimental.py drops from IDA land-bank written answers.

Status: sandbox prototype, NOT wired into pipeline.py. Reads only the disk
cache already populated by pq_answer_mine_experimental.py (no network for the
XML pass); writes to an isolated sandbox parquet.

Why it exists
--------------
parse_section() in pq_answer_mine_experimental.py only reads <p> children of
<speech> — it silently drops <table> elements. IDA land-bank disclosures are
exactly this shape: county/hectare grids inside <table><tr><td><p>. This
script re-walks the same cached AKN-XML sections and extracts the tables,
long-format (one row per cell), plus records any <a href> pointing at a
supportingDocumentation attachment (some answers, e.g. 52988/22, disclose via
an attached .docx instead of an inline table).

Input list of sections comes from data/_sandbox/pq_disclosures_full.parquet,
filtered for IDA land-bank questions (question_text matches IDA + a land-bank
keyword) — see _ida_land_refs(). 73 distinct sections matched 2026-08-01;
28 of them carry an inline <table>.

Run:
    python -m pipeline_sandbox.pq_disclosures.pq_ida_land_tables
"""

from __future__ import annotations

import io
import logging
import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path

import docx
import pandas as pd
import polars as pl

from pipeline_sandbox.pq_disclosures.pq_answer_mine_experimental import (
    _CACHE_DIR,
    _REF,
    _strip_ns,
    _text,
)
from services.http_engine import fetch_bytes
from services.logging_setup import setup_standalone_logging
from services.parquet_io import save_parquet

logger = logging.getLogger(__name__)

_FULL_CORPUS = Path("data/_sandbox/pq_disclosures_full.parquet")
_OUT_CELLS = Path("data/_sandbox/pq_ida_land_table_cells.parquet")
_OUT_ATTACH = Path("data/_sandbox/pq_ida_land_attachments.parquet")
_OUT_ATTACH_CELLS = Path("data/_sandbox/pq_ida_land_attachment_cells.parquet")
_ATTACH_CACHE = Path("c:/tmp/pq_answer_cache/attachments")

_Q_OPEN = re.compile(r"^\s*\d+\.\s+Deput(?:y|ies)\s+(.+?)\s+asked\b", re.I)
# Both boundaries required: a bare trailing \b (no leading \b) false-matched
# "spina bifida" (case-insensitive "...bIDA " satisfies "IDA\b") once tried.
# A loose "IDA + tabular anywhere" fallback was also tried and rejected — it
# pulled in unrelated IDA staffing/funding/Canadian-companies questions that
# happen to request tabular form. Keyword list tuned by manual sampling
# 2026-08-01 (n=12 on the full corpus); a bare "vacant" was tried and dropped
# — it matched "vacant positions" (staffing), so "vacant or derelict" is a
# phrase match instead. This is a recall-limited heuristic (Indicative band):
# it will miss land-bank questions phrased without any of these tokens.
_IDA_LAND_PAT = re.compile(
    r"(?i)\bIDA\b.*(hectare|land\s?bank|zoned?|acre|land\s+own|land\s+holding"
    r"|vacant\s+or\s+derelict|strategic\s+land|strategic\s+propert|land\s+portfolio)"
)
_KNOWN_GLUE_REFS = {"52988/22"}  # verified manually: "theIDA holds" — no space


def _ida_land_sections() -> list[dict]:
    """Distinct written-answer sections whose question text is IDA land-bank."""
    df = pl.scan_parquet(_FULL_CORPUS)
    hits = (
        df.filter(
            pl.col("question_text").str.contains(_IDA_LAND_PAT.pattern)
            | pl.col("question_ref").is_in(list(_KNOWN_GLUE_REFS))
        )
        .select(["question_ref", "date", "deputy", "department", "xml_uri"])
        .unique(subset=["xml_uri"])
        .collect()
    )
    return hits.to_dicts()


def _cache_path(xml_uri: str) -> Path:
    key = xml_uri.split("/debateRecord/", 1)[-1].replace("/", "_")
    return _CACHE_DIR / key


def _cell_rows(table_el: ET.Element) -> list[list[str]]:
    """<table><tr><td><p>...</p></td>...</tr>...</table> -> list of row-cell-text lists."""
    grid = []
    for tr in table_el:
        if _strip_ns(tr.tag) != "tr":
            continue
        cells = [_text(td) for td in tr if _strip_ns(td.tag) in ("td", "th")]
        grid.append(cells)
    return grid


def extract_tables_and_attachments(xml: str, meta: dict) -> tuple[list[dict], list[dict]]:
    """Walk one section's AKN-XML; return (table_cell_rows, attachment_rows).

    Mirrors parse_section()'s pending-question grouping so a table (or an
    attachment link) is attributed to the question_ref(s) it answers.
    """
    try:
        root = ET.fromstring(xml)
    except ET.ParseError as e:
        logger.warning("xml parse error %s: %s", meta.get("xml_uri"), e)
        return [], []

    cell_rows: list[dict] = []
    attach_rows: list[dict] = []
    pending: list[dict] = []

    for sec in root.iter():
        if _strip_ns(sec.tag) != "debateSection":
            continue
        for el in list(sec):
            tag = _strip_ns(el.tag)
            if tag == "question":
                txt = _text(el)
                m = _Q_OPEN.match(txt)
                for ref in _REF.findall(txt):
                    pending.append({"ref": ref, "deputy": (m.group(1) if m else None)})
            elif tag == "speech" and pending:
                table_idx = 0
                for child in el:
                    ctag = _strip_ns(child.tag)
                    if ctag == "table":
                        grid = _cell_rows(child)
                        if not grid:
                            continue
                        headers = grid[0]
                        for row_idx, row in enumerate(grid[1:]):
                            # A blank first cell marks a totals/footer row (seen
                            # in 6273/22's county table: "" | 1175.085 | ... —
                            # the column sums), not a data row — drop it.
                            if not row or not row[0].strip():
                                continue
                            for col_idx, val in enumerate(row):
                                col_name = headers[col_idx] if col_idx < len(headers) else None
                                for q in pending:
                                    cell_rows.append(
                                        {
                                            "question_ref": q["ref"],
                                            "deputy": q["deputy"],
                                            "date": meta["date"],
                                            "department": meta.get("department"),
                                            "table_index": table_idx,
                                            "row_index": row_idx,
                                            "col_index": col_idx,
                                            "col_name": col_name,
                                            "value": val,
                                            "source_xml_url": meta["xml_uri"],
                                        }
                                    )
                        table_idx += 1
                    else:
                        # supportingDocumentation attachment links live in <p><a href=...>
                        for a in child.iter():
                            if _strip_ns(a.tag) != "a":
                                continue
                            href = a.get("href") or ""
                            if "supportingDocumentation" not in href:
                                continue
                            for q in pending:
                                attach_rows.append(
                                    {
                                        "question_ref": q["ref"],
                                        "deputy": q["deputy"],
                                        "date": meta["date"],
                                        "department": meta.get("department"),
                                        "attachment_url": href,
                                        "link_text": _text(a),
                                        "source_xml_url": meta["xml_uri"],
                                    }
                                )
                pending = []
    return cell_rows, attach_rows


def _fetch_attachment(url: str) -> bytes | None:
    """Download a supportingDocumentation attachment, caching to disk (not
    populated by the XML-only pass, so this is a real network fetch)."""
    _ATTACH_CACHE.mkdir(parents=True, exist_ok=True)
    cache = _ATTACH_CACHE / url.rsplit("/", 1)[-1]
    if cache.exists():
        return cache.read_bytes()
    data = fetch_bytes(url, validate=lambda b: b[:2] == b"PK")  # docx/xlsx = zip
    if data:
        cache.write_bytes(data)
    return data


def parse_docx_tables(data: bytes, attach_meta: dict) -> list[dict]:
    """First table in a .docx attachment -> long-format cell rows."""
    doc = docx.Document(io.BytesIO(data))
    if not doc.tables:
        logger.warning("no tables in attachment %s", attach_meta["attachment_url"])
        return []
    table = doc.tables[0]
    grid = [[c.text.strip() for c in row.cells] for row in table.rows]
    if not grid:
        return []
    headers = grid[0]
    rows = []
    for row_idx, row in enumerate(grid[1:]):
        if not row or not row[0].strip():  # totals/footer row — see inline-table note above
            continue
        for col_idx, val in enumerate(row):
            col_name = headers[col_idx] if col_idx < len(headers) else None
            rows.append(
                {
                    "question_ref": attach_meta["question_ref"],
                    "deputy": attach_meta["deputy"],
                    "date": attach_meta["date"],
                    "row_index": row_idx,
                    "col_index": col_idx,
                    "col_name": col_name,
                    "value": val,
                    "attachment_url": attach_meta["attachment_url"],
                }
            )
    return rows


def main() -> int:
    setup_standalone_logging("pq_ida_land_tables")

    sections = _ida_land_sections()
    logger.info("IDA land-bank sections (distinct): %d", len(sections))

    all_cells: list[dict] = []
    all_attach: list[dict] = []
    missing_cache = 0
    for sec in sections:
        cache = _cache_path(sec["xml_uri"])
        if not cache.exists():
            missing_cache += 1
            continue
        cells, attach = extract_tables_and_attachments(cache.read_text(encoding="utf-8"), sec)
        all_cells.extend(cells)
        all_attach.extend(attach)

    n_sections_with_table = len({c["source_xml_url"] for c in all_cells})
    n_sections_with_attach = len({a["source_xml_url"] for a in all_attach})
    n_neither = len(sections) - missing_cache - len(
        {c["source_xml_url"] for c in all_cells} | {a["source_xml_url"] for a in all_attach}
    )
    logger.info("=" * 60)
    logger.info("sections total=%d missing_cache=%d", len(sections), missing_cache)
    logger.info("sections with >=1 inline table: %d (%d cells)", n_sections_with_table, len(all_cells))
    logger.info("sections with attachment link : %d", n_sections_with_attach)
    logger.info("sections with neither         : %d (narrative-only or unmatched shape)", n_neither)

    if all_cells:
        save_parquet(pd.DataFrame(all_cells), _OUT_CELLS)
        logger.info("wrote %s", _OUT_CELLS)
    if all_attach:
        save_parquet(pd.DataFrame(all_attach), _OUT_ATTACH)
        logger.info("wrote %s", _OUT_ATTACH)

    # Fetch + parse each attachment's own table (network — not cached by the
    # earlier XML-only pass).
    attach_cells: list[dict] = []
    for a in all_attach:
        data = _fetch_attachment(a["attachment_url"])
        if not data:
            logger.warning("attachment fetch failed: %s", a["attachment_url"])
            continue
        try:
            attach_cells.extend(parse_docx_tables(data, a))
        except Exception as e:
            logger.warning("attachment parse failed %s: %s", a["attachment_url"], e)
    if attach_cells:
        save_parquet(pd.DataFrame(attach_cells), _OUT_ATTACH_CELLS)
        n_rows = len({(c["question_ref"], c["row_index"]) for c in attach_cells})
        logger.info("wrote %s (%d attachment table rows, %d cells)", _OUT_ATTACH_CELLS, n_rows, len(attach_cells))
    return 0


if __name__ == "__main__":
    sys.exit(main())
