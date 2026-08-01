"""
pq_attachment_harvest.py — EXPERIMENTAL: fetch and parse the ATTACHED files
(docx/xlsx/xls/pdf) that ministers use for their biggest written-answer
disclosures.

Status: sandbox prototype, NOT wired into pipeline.py. Writes to isolated
sandbox parquets. Method + traps: memory reference_pq_attachment_harvest_2026_08_01.

Why
---
When an answer's data is too big to inline, it arrives as a linked file. The
census found 3,473 of them (docx 1,323 / pdf 1,098 / xlsx 990 / doc 27 / xls 13
/ png 13) and, before this script, exactly ONE had been parsed. An answer that
inlines a table gives 5-25 rows; the one attachment parsed so far gave 317. This
is where itemised registers live.

Output is LONG format (one row per cell) so heterogeneous tables share a schema:
    source_ref | date | department | attachment_url | sheet_or_table
    | row_index | col_index | col_name | value

Formats
-------
docx  -> python-docx, every table in the document
xlsx  -> openpyxl (read_only), every sheet
xls   -> xlrd via pandas
pdf   -> fitz (PyMuPDF) native table detection, `page.find_tables()`, with page
         text as the fallback when no table is found. NOT camelot: user
         directive 2026-08-01, and independently ghostscript is not installed
         on this box (`gs: command not found`), which camelot's lattice mode
         requires. camelot stays confined to the two AFS extractors' isolated
         venv ($AFS_CAMELOT_VENV) — don't uv-add it here.

Bounded + resumable: --limit caps downloads, every file is disk-cached, so a
killed run resumes for free and reruns cost no network.

Run (smoke, 40 files):
    python -m pipeline_sandbox.pq_disclosures.pq_attachment_harvest --limit 40
Full channel:
    python -m pipeline_sandbox.pq_disclosures.pq_attachment_harvest --limit 0 --workers 6
Filter to a topic:
    ... --filter "waiting|hospital"     (regex on department/section_title/link_text)
"""

from __future__ import annotations

import services.runtime_env  # noqa: F401  # MUST be first: caps BLAS threads

import argparse
import io
import logging
import re
import sys
import time
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import polars as pl

from services.http_engine import fetch_bytes
from services.logging_setup import setup_standalone_logging
from services.parquet_io import save_parquet

logger = logging.getLogger(__name__)

_CENSUS = Path("data/_sandbox/pq_attachment_census.parquet")
_CACHE = Path("c:/tmp/pq_answer_cache/attachments")
_OUT_CELLS = Path("data/_sandbox/pq_attachment_cells.parquet")
_OUT_INDEX = Path("data/_sandbox/pq_attachment_index.parquet")

_ZIP_EXT = {"docx", "xlsx"}
_MAX_ROWS_PER_TABLE = 5_000   # a runaway sheet shouldn't blow the parquet


def _validator(ext: str):
    """Content sniff — catches a WAF/error page returned as HTTP 200."""
    if ext in _ZIP_EXT:
        return lambda b: b[:2] == b"PK"
    if ext == "pdf":
        return lambda b: b[:4] == b"%PDF"
    return lambda b: len(b) > 0


def _normalise_url(url: str) -> str:
    """Repair hrefs that are malformed IN THE SOURCE XML.

    74 of 3,473 (2.1%) don't parse as-is: 66 carry a single slash after the
    scheme ("https:/data.oireachtas.ie/...") and 8 omit the scheme entirely
    ("data.oireachtas.ie/..."). urllib quoting preserves both faithfully, so
    the requests leg 404s and only the curl fallback has any chance. Repair
    before fetching rather than losing them.
    """
    url = url.strip()
    # Stray characters the real filename does not contain: backslash-escaped
    # underscores ("2025-03-19\_pq1446…", 6 cases) and interior spaces
    # ("…/ 2020-07-30_pq93…", "…_en .docx", 6 cases). Quoting turns these into
    # %5C / %20 and the fetch 404s. Both are safe to delete outright.
    url = url.replace("\\", "").replace(" ", "")
    url = re.sub(r"^(https?):/(?!/)", r"\1://", url)
    if url.startswith("//"):
        url = "https:" + url
    elif not url.startswith(("http://", "https://")):
        url = "https://" + url.lstrip("/")
    # Collapse doubled slashes in the PATH only ("…/supportingDocumentation//file.docx"),
    # never in the "https://" scheme separator.
    scheme, rest = url.split("://", 1)
    return f"{scheme}://{re.sub(r'/{2,}', '/', rest)}"


def _fetch(url: str, ext: str) -> bytes | None:
    url = _normalise_url(url)
    _CACHE.mkdir(parents=True, exist_ok=True)
    # Filenames are NOT unique across answers, so key on the full path tail.
    key = re.sub(r"[^A-Za-z0-9._-]", "_", url.split("supportingDocumentation/", 1)[-1])
    cached = _CACHE / key
    if cached.exists():
        return cached.read_bytes()
    data = fetch_bytes(url, validate=_validator(ext))
    if data:
        cached.write_bytes(data)
        time.sleep(0.1)  # stay a polite citizen; cache-hits skip this entirely
    return data


def _emit(grid: list[list[str]], where: str) -> list[dict]:
    """A 2-D grid -> long-format cell rows, applying the shared table traps."""
    grid = [r for r in grid if any((c or "").strip() for c in r)]
    if not grid:
        return []
    header = [(c or "").strip() for c in grid[0]]
    out = []
    for row_index, row in enumerate(grid[1 : 1 + _MAX_ROWS_PER_TABLE]):
        # Blank first cell = totals/footer row, not data (see memory card).
        if not row or not (row[0] or "").strip():
            continue
        for col_index, val in enumerate(row):
            # Skip EMPTY cells. openpyxl returns a sheet's whole used rectangle,
            # which stray formatting inflates enormously: one OPW attachment
            # emitted 1,326,861 cells of which 959 were populated (0.07%), and
            # 65.2% of a first full harvest was empty padding. row_index /
            # col_index still carry the grid position, so nothing is lost —
            # and ranking attachments by cell count stops being meaningless.
            text = (val or "").strip()
            if not text:
                continue
            out.append(
                {
                    "sheet_or_table": where,
                    "row_index": row_index,
                    "col_index": col_index,
                    "col_name": header[col_index] if col_index < len(header) else None,
                    "value": text,
                }
            )
    return out


def parse_docx(data: bytes) -> list[dict]:
    import docx

    doc = docx.Document(io.BytesIO(data))
    rows: list[dict] = []
    for i, t in enumerate(doc.tables):
        rows.extend(_emit([[c.text for c in r.cells] for r in t.rows], f"table_{i}"))
    return rows


def parse_xlsx(data: bytes) -> list[dict]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows: list[dict] = []
    for ws in wb.worksheets:
        grid = [
            ["" if c is None else str(c) for c in r]
            for r in ws.iter_rows(max_row=_MAX_ROWS_PER_TABLE + 1, values_only=True)
        ]
        rows.extend(_emit(grid, ws.title))
    wb.close()
    return rows


def parse_xls(data: bytes) -> list[dict]:
    import pandas as pd

    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None, dtype=str)
    rows: list[dict] = []
    for name, sdf in sheets.items():
        rows.extend(_emit(sdf.fillna("").astype(str).values.tolist(), str(name)))
    return rows


def parse_pdf(data: bytes) -> list[dict]:
    """Extract PDF TABLES with fitz (PyMuPDF) — no camelot, per user direction
    2026-08-01.

    PyMuPDF has had native table detection since 1.23 (`page.find_tables()`,
    confirmed present on 1.27.2.3 here), so this needs no isolated venv and no
    subprocess hop. Falls back to page text only for pages where no table is
    detected, so a text-only PDF still yields something.
    """
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    rows: list[dict] = []
    try:
        for pno, page in enumerate(doc):
            found = 0
            try:
                for ti, table in enumerate(page.find_tables().tables):
                    grid = [["" if c is None else str(c) for c in r] for r in table.extract()]
                    emitted = _emit(grid, f"page_{pno}_table_{ti}")
                    rows.extend(emitted)
                    found += len(emitted)
            except Exception as e:  # detection can fail on odd page trees
                logger.debug("find_tables failed p%d: %s", pno, e)
            if not found:
                text = page.get_text().strip()
                if text:
                    rows.append(
                        {
                            "sheet_or_table": f"page_{pno}_text",
                            "row_index": 0,
                            "col_index": 0,
                            "col_name": "page_text",
                            "value": text[:20_000],
                        }
                    )
    finally:
        doc.close()
    return rows


_PARSERS = {"docx": parse_docx, "xlsx": parse_xlsx, "xls": parse_xls, "pdf": parse_pdf}


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--limit", type=int, default=40, help="max attachments; 0 = all")
    ap.add_argument("--workers", type=int, default=4, help="concurrent downloads (keep <= 20)")
    ap.add_argument("--filter", default=None, help="regex on department/section_title/link_text")
    ap.add_argument("--ext", default="docx,xlsx,xls", help="comma list of extensions to parse")
    args = ap.parse_args(argv)

    setup_standalone_logging("pq_attachment_harvest")

    want = {e.strip() for e in args.ext.split(",") if e.strip()}
    df = pl.read_parquet(_CENSUS).filter(pl.col("ext").is_in(list(want)))
    if args.filter:
        pat = f"(?i){args.filter}"
        df = df.filter(
            pl.col("department").fill_null("").str.contains(pat)
            | pl.col("section_title").fill_null("").str.contains(pat)
            | pl.col("link_text").fill_null("").str.contains(pat)
        )
    # One row per distinct URL: the same file is linked from several answers.
    df = df.unique(subset=["attachment_url"]).sort("date", descending=True)
    if args.limit:
        df = df.head(args.limit)
    targets = df.to_dicts()
    logger.info("attachments to harvest: %d (ext=%s, filter=%s)", len(targets), sorted(want), args.filter)

    cells: list[dict] = []
    index: list[dict] = []

    def _one(rec: dict) -> None:
        ext = rec["ext"]
        data = _fetch(rec["attachment_url"], ext)
        status, n = "ok", 0
        if not data:
            status = "fetch_failed"
        else:
            try:
                parsed = _PARSERS[ext](data)
                n = len(parsed)
                for c in parsed:
                    cells.append({**c, **{k: rec[k] for k in ("date", "department", "section_title", "attachment_url")}})
                if n == 0:
                    status = "no_tables"
                elif ext == "pdf" and all(c.get("col_name") == "page_text" for c in parsed):
                    # fitz found no table on any page — text captured only.
                    status = "pdf_text_only"
            except Exception as e:  # one bad file must not kill the run
                status = f"parse_error: {type(e).__name__}"
                logger.warning("parse failed %s: %s", rec["attachment_url"], e)
        index.append({**{k: rec[k] for k in ("date", "department", "section_title", "link_text", "attachment_url", "ext")},
                      "status": status, "n_cells": n})

    with ThreadPoolExecutor(max_workers=max(1, args.workers)) as ex:
        list(ex.map(_one, targets))

    if index:
        idx = pl.DataFrame(index)
        save_parquet(idx, _OUT_INDEX)
        logger.info("=" * 60)
        logger.info("attempted %d", idx.height)
        for r in idx.group_by("status").len().sort("len", descending=True).iter_rows():
            logger.info("  %-28s %d", r[0], r[1])
        ok = idx.filter(pl.col("n_cells") > 0)
        if ok.height:
            logger.info("richest attachments:")
            for r in ok.sort("n_cells", descending=True).head(8).iter_rows(named=True):
                logger.info("  %5d cells | %-14s | %s", r["n_cells"], (r["department"] or "")[:14],
                            (r["link_text"] or r["attachment_url"].rsplit("/", 1)[-1])[:60])
    if cells:
        save_parquet(pl.DataFrame(cells), _OUT_CELLS)
        logger.info("wrote %s (%d cells)", _OUT_CELLS, len(cells))
    return 0


if __name__ == "__main__":
    sys.exit(main())
